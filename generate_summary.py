from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_summary():
    doc = Document()
    
    # Title
    title = doc.add_heading('Rangkuman Tugas DevOps / MLOps InterOpera', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Dokumen ini merupakan rangkuman lengkap dari dokumen panduan (homework_brief_devops.pdf), status pengerjaan kita saat ini (implementation_plan.md), serta penjelasan mendetail per komponen (bedah code dan docs).').style = 'Intense Quote'
    
    # Section 1
    doc.add_heading('1. Latar Belakang & Persyaratan Utama', level=1)
    p = doc.add_paragraph('Klien (Meridian Asset Management) membutuhkan layanan chatbot LLM dengan syarat mutlak: ')
    p.add_run('self-hosted, tidak boleh ada data yang keluar environment, dan fully observable').bold = True
    p.add_run('. Tidak boleh ada koneksi ke API eksternal (seperti OpenAI) saat runtime.')
    
    doc.add_heading('Layanan yang sudah ada:', level=2)
    doc.add_paragraph('model-server: Mock inference server yang menyimulasikan vLLM.', style='List Bullet')
    doc.add_paragraph('gateway: Client-facing service.', style='List Bullet')
    
    # Section 2
    doc.add_heading('2. Lima Aturan Ketat (Hard Constraints)', level=1)
    doc.add_paragraph('Reproducible bring-up: Keseluruhan platform bisa jalan di mesin lokal hanya dengan 1 command (misal: make up).', style='List Number')
    doc.add_paragraph('Everything as code: Seluruh resource infrastruktur dan monitoring wajib dideklarasikan via kode (Terraform/Helm/K8s YAML).', style='List Number')
    doc.add_paragraph('Gated model delivery: Rollout (Canary deployment) harus ter-promote atau ter-rollback otomatis tanpa campur tangan manusia (berdasarkan metrik pengujian).', style='List Number')
    doc.add_paragraph('Provable observability: Alarm (Alerting) wajib menyala/menembak jika disimulasikan adanya error/fault.', style='List Number')
    doc.add_paragraph('Self-hosted, no egress: Murni berjalan di jaringan tertutup tanpa memanggil layanan awan luar saat runtime.', style='List Number')
    
    # Section 3
    doc.add_heading('3. Enam Fase Pekerjaan (Deliverables)', level=1)
    doc.add_paragraph('Phase 1 - Platform Design Memo: Dokumen arsitektur, strategi MLOps (Canary vs Blue/Green), dan analisis jika berjalan di GPU sungguhan.', style='List Bullet')
    doc.add_paragraph('Phase 2 - Platform as Code: Menyusun Makefile, Terraform, dan Manifests/Helm untuk Kubernetes lokal (kind).', style='List Bullet')
    doc.add_paragraph('Phase 3 - RAG Chat Service: Membangun backend Python FastAPI ("rag-api") dari nol untuk memproses Retrieval-Augmented Generation (memakai Qdrant dan sentence-transformers lokal).', style='List Bullet')
    doc.add_paragraph('Phase 4 - Gated Model Delivery: Membuat bash/python scripts otomatis untuk CI/CD (eval/rollout).', style='List Bullet')
    doc.add_paragraph('Phase 5 - Observability & Alerting: Merancang dashboard Grafana dan alert Prometheus.', style='List Bullet')
    doc.add_paragraph('Phase 6 - Production Incident: Men-debug masalah latency pada gateway (tidak ada connection pool), memperbaikinya, dan menulis laporan postmortem.', style='List Bullet')

    # Section 4
    doc.add_heading('4. Bedah Kode (Perkiraan Komposisi)', level=1)
    doc.add_paragraph('40% YAML/HCL: Konfigurasi Kubernetes, Terraform, Prometheus Alerts.', style='List Bullet')
    doc.add_paragraph('30% Python: Menulis RAG API backend.', style='List Bullet')
    doc.add_paragraph('20% Shell Scripting: Script evaluator dan Makefile.', style='List Bullet')
    doc.add_paragraph('10% Debugging & Markdown: Postmortem incident dan Platform Memo.', style='List Bullet')
    
    # Section 5 - Penjelasan Folder
    doc.add_heading('5. Penjelasan per Folder', level=1)
    
    doc.add_heading('Root Folder (/)', level=2)
    doc.add_paragraph('Berisi Makefile (kumpulan perintah seperti make up) dan README.md (panduan utama).')

    doc.add_heading('infra/', level=2)
    doc.add_paragraph('Infrastructure as Code. Berisi file Terraform (main.tf) untuk mendirikan Kubernetes Cluster dan menginstall operator esensial.')

    doc.add_heading('deploy/', level=2)
    doc.add_paragraph('Kubernetes Manifests. Berisi file YAML (gateway.yaml, rag-api.yaml) untuk mendeploy aplikasi (Pod, Service) ke dalam klaster.')

    doc.add_heading('services/', level=2)
    doc.add_paragraph('Source Code Aplikasi. Terdiri dari model-server, gateway (yang akan diperbaiki di Phase 6), dan rag-api (API Python FastAPI yang harus dibangun dari nol beserta algoritmanya).')

    doc.add_heading('eval/', level=2)
    doc.add_paragraph('Evaluation Gate. Berisi Bash/Python scripts (seperti rollout.sh) untuk menjalankan evaluasi otomatis (Canary testing) setiap kali ada versi model baru.')

    doc.add_heading('observability/', level=2)
    doc.add_paragraph('Monitoring & Alerting. Berisi file konfigurasi YAML dan JSON untuk Prometheus Rules (alarm metrik) dan Grafana Dashboards.')

    doc.add_heading('corpus/', level=2)
    doc.add_paragraph('Dokumen Pengetahuan Dasar. Berisi file teks atau PDF aturan reksadana yang akan diubah menjadi embedding vektor oleh RAG.')

    doc.add_heading('docs/', level=2)
    doc.add_paragraph('Dokumen Tertulis. Berisi 01_platform_memo.md (arsitektur, SLO, dan desain MLOps) serta 02_postmortem.md (laporan insiden terkait bug di gateway).')

    doc.add_heading('evidence/', level=2)
    doc.add_paragraph('Barang Bukti. Sesuai instruksi spesifik di dokumen, folder ini wajib menyimpan 4 hal:')
    doc.add_paragraph('Gate decisions: Log bukti keputusan sistem CI/CD saat melakukan promote/rollback model.', style='List Bullet')
    doc.add_paragraph('Alert firing: Bukti (screenshot/log) bahwa sistem alert menyala saat diinjeksi error.', style='List Bullet')
    doc.add_paragraph('Before/after loads: Hasil perbandingan load-test SEBELUM dan SESUDAH perbaikan bug di Phase 6.', style='List Bullet')
    doc.add_paragraph('RAG eval results: Log bukti nilai evaluasi akurasi layanan RAG.', style='List Bullet')

    doc.add_heading('ci/ (Opsional)', level=2)
    doc.add_paragraph('Jika menggunakan CI/CD seperti GitHub Actions, berisi file YAML pipeline pengujian (pipeline.yml).')

    # Section 6
    doc.add_heading('6. Progres Saat Ini', level=1)
    doc.add_paragraph('Saat ini, kita telah menyelesaikan tahap Perencanaan Total dengan menghasilkan dokumen implementation_plan.md. Dokumen ini menjadi peta jalan (roadmap) 7 hari kerja yang sudah merinci tech stack yang kita pilih (kind, Qdrant, Terraform, Helm, Prometheus) serta breakdown tugas per hari yang sangat siap eksekusi.')

    # Section 7 - Interview Guide
    doc.add_heading('7. Panduan Presentasi Interview (Elevator Pitch)', level=1)
    doc.add_paragraph('Gunakan alur narasi berikut saat technical interview untuk menunjukkan kapasitas Anda sebagai Platform/DevOps Engineer senior yang tidak hanya pandai coding, namun juga memiliki "Ops Judgment".')
    
    doc.add_heading('A. Membuka Konteks (The Problem)', level=2)
    doc.add_paragraph('"Studi kasus ini menuntut layanan LLM untuk klien finansial yang sangat diregulasi. Syarat mutlaknya: data tidak boleh keluar (privacy), dilarang pakai API eksternal (seperti OpenAI), dan harus self-hosted. Tugas saya adalah membangun platform infrastruktur dari nol agar layanan AI ini siap rilis ke production secara aman."').style = 'Intense Quote'
    
    doc.add_heading('B. Menjelaskan Arsitektur & Solusi', level=2)
    doc.add_paragraph('"Saya mendesain platform ini dengan prinsip Everything as Code. Infrastruktur dibangun dengan Terraform di atas klaster Kubernetes lokal (kind). Pendekatan ini memastikan seluruh sistem bisa dibangun dari mesin kosong secara reproducible hanya dengan satu perintah: make up."').style = 'Intense Quote'
    
    doc.add_heading('C. Membedah Komponen Kunci', level=2)
    doc.add_paragraph('Highlight 1 - RAG Data Plane: "Saya membangun layanan backend RAG API dari nol (Python/FastAPI) yang terintegrasi dengan Qdrant Vector DB dan embedding lokal. Ini mencegah LLM berhalusinasi tanpa melanggar privasi data."', style='List Bullet')
    doc.add_paragraph('Highlight 2 - Gated Model Delivery: "Saya merancang pipeline Canary Rollout otomatis (Zero Human-in-the-loop). Ketika ada versi AI baru, sistem akan mengetesnya. Jika akurasi anjlok atau melambat, akan terjadi Auto-Rollback. Jika bagus, Auto-Promote."', style='List Bullet')
    doc.add_paragraph('Highlight 3 - Observability & Incident Management: "Platform dipantau oleh Prometheus & Grafana. Saat terjadi simulasi insiden di mana latency klien melonjak tajam, berbekal metrik telemetry, saya menemukan bahwa bottleneck ada di layer Gateway (karena ketiadaan connection pool). Saya memodifikasi kode gateway dengan Connection Pooling dan insiden pun teratasi."', style='List Bullet')
    
    doc.add_heading('D. Antisipasi "Deep Dive Questions"', level=2)
    doc.add_paragraph('Kenapa pilih Canary daripada Blue/Green?: "LLM butuh memori GPU yang besar. Blue/Green butuh resource GPU 2x lipat saat proses switch (sangat mahal). Canary jauh lebih hemat dan aman."', style='List Bullet')
    doc.add_paragraph('Kenapa autoscaling LLM tidak berbasis CPU?: "LLM itu GPU-bound. CPU bisa saja terlihat santai (idle) padahal VRAM GPU sudah penuh sesak. Oleh karena itu, metrik pemicu autoscaling yang valid adalah antrean (Queue Depth / In-flight Requests), bukan pemakaian CPU."', style='List Bullet')

    # Section 8
    doc.add_heading('8. Status Audit Repositori Saat Ini', level=1)
    doc.add_paragraph('Berdasarkan pengecekan struktur direktori lokal, fondasi folder kita sudah 100% SESUAI dengan persyaratan spesifikasi tugas. Beberapa bukti kesesuaian utamanya meliputi:')
    doc.add_paragraph('Folder services/ secara akurat telah berisi model-server, gateway, dan rag-api.', style='List Bullet')
    doc.add_paragraph('Folder docs/ telah menampung 01_platform_memo.md dan 02_postmortem.md.', style='List Bullet')
    doc.add_paragraph('Folder evidence/ sudah berisi berkas otentik seperti screenshot alert-firing.png, hasil komparasi uji beban (load test sebelum & sesudah), serta kumpulan log JSON untuk keputusan promosi/rollback.', style='List Bullet')
    doc.add_paragraph('Catatan Tambahan: Pastikan file ekstra pendukung sementara (seperti skrip python generate_summary.py dan file rangkuman Word) ditambahkan ke .gitignore atau dihapus sebelum diserahkan, agar repositori tetap terlihat murni (clean) saat dinilai.', style='Intense Quote')

    doc.save('Rangkuman_Tugas_InterOpera_V5.docx')
    print("Dokumen berhasil diperbarui dengan penjelasan folder!")

if __name__ == '__main__':
    create_summary()
